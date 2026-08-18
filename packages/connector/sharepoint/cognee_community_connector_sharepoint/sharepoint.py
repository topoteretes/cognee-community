"""SharePoint / OneDrive connector for cognee — a ``dlt`` source over a document library.

Sync SharePoint document libraries (PDF, Word, Excel, and plain-text files) into
cognee, incrementally and with forget-on-deletion.  Built entirely on the
existing DLT ingestion subsystem; the resource produced here is handed directly
to :func:`cognee.remember`::

    import cognee
    from cognee_community_connector_sharepoint import sharepoint_source

    await cognee.remember(
        sharepoint_source(site_id="contoso.sharepoint.com,<guid>,<guid>"),
        dataset_name="my_sharepoint",
        primary_key="id",
        write_disposition="merge",   # incremental upsert by item id
    )

Design
------
* **Auth** — Microsoft Graph app-only (client credentials) via ``msal``.  This
  is the non-interactive default, mirroring how the Google Drive connector
  defaults to a service account.  The Entra ID app registration is shared with
  the planned Teams and Outlook connectors, so the credentials are read from
  ``MICROSOFT_TENANT_ID`` / ``MICROSOFT_CLIENT_ID`` / ``MICROSOFT_CLIENT_SECRET``
  rather than ``SHAREPOINT_*``.
* **Primary key** — ``"{drive_id}:{item_id}"``.  DriveItem ids are unique within
  a drive but not across drives, so the drive id disambiguates them when several
  libraries are synced into one dataset.
* **Incremental cursor** — the Graph delta link, one per drive, persisted in
  dlt's per-resource state.  The first run enumerates the library with a
  tokenless ``delta`` call and stores the returned ``@odata.deltaLink``; later
  runs follow that link and receive only what changed.  Graph documents ``delta``
  as the only way to enumerate a drive completely, so it is used for the initial
  read as well rather than a separate listing pass.
* **Forget-on-delete** — items carrying Graph's ``deleted`` facet are emitted
  with the ``_deleted`` hard-delete marker; dlt drops them from its destination
  on ``merge`` and cognee's existing ``orphan_cleanup`` purges them from the
  graph, vector, and relational stores.
* **Content** — PDFs are parsed with the core ``pypdf`` dependency, ``.docx``
  and ``.xlsx`` with ``python-docx`` and ``openpyxl``, and ``.txt`` / ``.md`` /
  ``.csv`` are decoded as-is.  Anything else — including ``.pptx`` and the legacy
  binary ``.doc`` / ``.xls`` formats — is skipped with a warning naming the file,
  as is any file above ``max_file_size_mb``.  A file that cannot be parsed is
  skipped the same way rather than failing the sync.

Limitations
-----------
* Scope is a site and its document libraries.  Folder-level scoping is not
  supported: delta responses omit ``parentReference.path``, so a subtree cannot
  be filtered client-side, and Graph's remarks note that renaming a folder does
  not re-report its descendants.
* Graph can invalidate a delta link at any time (no lifetime is published for
  driveItem tokens), answering ``HTTP 410`` with ``resyncChangesApplyDifferences``
  or ``resyncChangesUploadDifferences``.  The connector drops the link and
  re-enumerates the library from scratch.  Live files are re-emitted and merged
  idempotently, but a file deleted *during* the invalidation window is never
  reported as deleted and so survives in memory until it is removed again.
* Graph replays: the same item can appear several times in one delta feed.  Only
  the last occurrence for an id is used, per Graph's guidance.
* Server-side conversion (``?format=pdf``) would cover ``.pptx`` and the legacy
  Office formats without extra parsers, but its least-privileged application
  permission is ``Files.ReadWrite.All``.  This connector stays read-only instead.
"""

import io
import os
import time
from dataclasses import dataclass
from typing import Any

from cognee.shared.logging_utils import get_logger
from cognee.tasks.ingestion.dlt_utils import DOCUMENT_SOURCE_ATTR

logger = get_logger("sharepoint_connector")

GRAPH_BASE_URL = "https://graph.microsoft.com/v1.0"
# App-only tokens carry the permissions granted to the app registration itself.
GRAPH_DEFAULT_SCOPE = "https://graph.microsoft.com/.default"

_MAX_RETRIES = 5
_RETRY_STATUSES = frozenset({429, 500, 502, 503, 504})

PLAIN_TEXT_EXTENSIONS = frozenset({".txt", ".md", ".csv"})


@dataclass(frozen=True)
class _SharePointConfig:
    site_id: str | None
    drive_ids: tuple[str, ...]
    max_file_size_mb: int


# ---------------------------------------------------------------------------
# Auth / session construction
# ---------------------------------------------------------------------------
def build_graph_session(
    *,
    tenant_id: str | None = None,
    client_id: str | None = None,
    client_secret: str | None = None,
) -> Any:
    """Build a ``requests`` session that signs every call with an app-only token.

    ``msal`` and ``requests`` are imported lazily so they remain an optional
    dependency (``pip install "cognee[sharepoint]"``). Building the session
    contacts Entra to resolve the tenant, so bad credentials fail here rather
    than on the first Graph call.
    """
    try:
        import msal
        import requests
    except ImportError as exc:  # pragma: no cover - depends on optional extra
        raise ImportError(
            'The SharePoint connector requires the "sharepoint" extra. '
            'Install it with: pip install "cognee[sharepoint]" '
            "(provides msal and requests)."
        ) from exc

    resolved_tenant_id = tenant_id or os.getenv("MICROSOFT_TENANT_ID")
    resolved_client_id = client_id or os.getenv("MICROSOFT_CLIENT_ID")
    resolved_client_secret = client_secret or os.getenv("MICROSOFT_CLIENT_SECRET")
    if not (resolved_tenant_id and resolved_client_id and resolved_client_secret):
        raise ValueError(
            "Microsoft Graph app-only credentials are required: pass tenant_id, "
            "client_id and client_secret, or set MICROSOFT_TENANT_ID, "
            "MICROSOFT_CLIENT_ID and MICROSOFT_CLIENT_SECRET."
        )

    try:
        app = msal.ConfidentialClientApplication(
            resolved_client_id,
            authority=f"https://login.microsoftonline.com/{resolved_tenant_id}",
            client_credential=resolved_client_secret,
        )
    except ValueError as exc:
        # msal resolves the authority over the network here, so a wrong tenant id
        # fails at construction rather than on the first Graph call.
        raise ValueError(
            f"Microsoft Graph authentication failed: could not resolve tenant "
            f"{resolved_tenant_id!r}. Check MICROSOFT_TENANT_ID. ({exc})"
        ) from exc

    def sign(request):
        request.headers["Authorization"] = f"Bearer {acquire_token(app)}"
        return request

    session = requests.Session()
    session.auth = sign
    session.headers.update({"Accept": "application/json"})
    return session


def acquire_token(app: Any) -> str:
    """Return an app-only access token, surfacing Entra's own error text."""
    # msal caches the token and only calls Entra again when it nears expiry, so
    # this is safe to call per request and keeps long syncs from expiring.
    result = app.acquire_token_for_client(scopes=[GRAPH_DEFAULT_SCOPE])
    token = result.get("access_token")
    if not token:
        detail = result.get("error_description") or result.get("error") or "unknown error"
        raise ValueError(f"Microsoft Graph authentication failed: {detail}")
    return token


# ---------------------------------------------------------------------------
# Graph HTTP helpers
# ---------------------------------------------------------------------------
def _request(session: Any, url: str, params: dict | None = None) -> Any:
    """GET a Graph URL, retrying throttled and transient responses."""
    for attempt in range(_MAX_RETRIES):
        response = session.get(url, params=params or {})
        if response.status_code in _RETRY_STATUSES and attempt < _MAX_RETRIES - 1:
            delay = _retry_after(response, attempt)
            logger.warning(
                "SharePoint: Graph returned %s — retrying in %.1fs (%d/%d).",
                response.status_code,
                delay,
                attempt + 1,
                _MAX_RETRIES,
            )
            time.sleep(delay)
            continue
        response.raise_for_status()
        return response


def _retry_after(response: Any, attempt: int) -> float:
    """Seconds to wait before retrying: the Retry-After header, else backoff."""
    header = (getattr(response, "headers", None) or {}).get("Retry-After")
    try:
        return float(header)
    except (TypeError, ValueError):
        return float(2**attempt)


def _api_get(session: Any, url: str, params: dict | None = None) -> dict:
    return _request(session, url, params).json()


def _is_resync_required(exc: Exception) -> bool:
    """True when Graph invalidated the delta link and wants a full re-enumeration."""
    # Trust only the structured status: a Graph error body embeds the request
    # URL, which can contain "410" by coincidence.
    return getattr(getattr(exc, "response", None), "status_code", None) == 410


# ---------------------------------------------------------------------------
# Content extraction (extension dispatch)
# ---------------------------------------------------------------------------
def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data), strict=False)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(", ".join(cells))
    return "\n".join(parts)


def _extract_xlsx_text(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        lines: list[str] = []
        for sheet in workbook.worksheets:
            lines.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(value) for value in row if value is not None]
                if cells:
                    lines.append(", ".join(cells))
        return "\n".join(lines)
    finally:
        workbook.close()


_EXTRACTORS = {
    ".pdf": _extract_pdf_text,
    ".docx": _extract_docx_text,
    ".xlsx": _extract_xlsx_text,
}


def is_supported_file(name: str) -> bool:
    extension = os.path.splitext(name)[1].lower()
    return extension in _EXTRACTORS or extension in PLAIN_TEXT_EXTENSIONS


def extract_file_content(session: Any, drive_id: str, item_id: str, name: str) -> str | None:
    """Return extracted text for a drive item, or None to skip it.

    A file that cannot be downloaded or parsed is logged and skipped rather than
    aborting the whole library sync. Auth and connectivity failures surface from
    the delta calls instead.
    """
    extension = os.path.splitext(name)[1].lower()
    try:
        data = _request(
            session, f"{GRAPH_BASE_URL}/drives/{drive_id}/items/{item_id}/content"
        ).content
        if extension in _EXTRACTORS:
            return _EXTRACTORS[extension](data)
        if extension in PLAIN_TEXT_EXTENSIONS:
            return data.decode("utf-8", errors="replace")
    except Exception as exc:
        logger.warning(
            "Skipping SharePoint file '%s' (%s): content extraction failed: %s", name, item_id, exc
        )
        return None

    # Unsupported extension. _item_to_row guards with is_supported_file and logs
    # the skip, so don't log twice.
    return None


# ---------------------------------------------------------------------------
# Sync state machine (pure given a session + state dict — unit-testable)
# ---------------------------------------------------------------------------
def _resolve_drive_ids(session: Any, config: _SharePointConfig) -> list[str]:
    if config.drive_ids:
        return list(config.drive_ids)

    try:
        payload = _api_get(session, f"{GRAPH_BASE_URL}/sites/{config.site_id}/drives")
    except Exception as exc:
        raise RuntimeError(
            f"SharePoint: failed to list document libraries for site '{config.site_id}': {exc}"
        ) from exc
    return [drive["id"] for drive in payload.get("value", []) if drive.get("id")]


def _collect_delta(session: Any, url: str) -> tuple[dict[str, dict], str | None]:
    """Follow a delta feed to its end, returning the latest state per item id."""
    items: dict[str, dict] = {}
    delta_link: str | None = None

    while url:
        payload = _api_get(session, url)
        for item in payload.get("value", []):
            item_id = item.get("id")
            if item_id:
                # Graph may replay an item within one feed; last occurrence wins.
                items[item_id] = item
        delta_link = payload.get("@odata.deltaLink") or delta_link
        url = payload.get("@odata.nextLink")

    return items, delta_link


def _item_to_row(session: Any, config: _SharePointConfig, drive_id: str, item: dict) -> dict | None:
    item_id = item["id"]
    name = item.get("name") or ""

    if "file" not in item:
        return None

    if not is_supported_file(name):
        logger.warning(
            "Skipping unsupported SharePoint file '%s' (%s): extension is not one of %s.",
            name,
            item_id,
            sorted(set(_EXTRACTORS) | PLAIN_TEXT_EXTENSIONS),
        )
        return None

    size = item.get("size")
    if size and int(size) > config.max_file_size_mb * 1024 * 1024:
        logger.warning(
            "Skipping SharePoint file '%s' (%s): size exceeds max_file_size_mb=%d.",
            name,
            item_id,
            config.max_file_size_mb,
        )
        return None

    content = extract_file_content(session, drive_id, item_id, name)
    if content is None or not content.strip():
        return None

    return {
        "id": _row_id(drive_id, item_id),
        "title": name,
        "content": content,
        "url": item.get("webUrl"),
        "_deleted": False,
    }


def _row_id(drive_id: str, item_id: str) -> str:
    return f"{drive_id}:{item_id}"


def _sync_drive(session: Any, config: _SharePointConfig, drive_id: str, delta_links: dict):
    """Yield rows for one document library and advance its delta link."""
    initial_url = f"{GRAPH_BASE_URL}/drives/{drive_id}/root/delta"
    stored_link = delta_links.get(drive_id)

    try:
        items, delta_link = _collect_delta(session, stored_link or initial_url)
    except Exception as exc:
        if not (stored_link and _is_resync_required(exc)):
            raise
        logger.warning(
            "SharePoint: delta link for drive %s was invalidated by Graph; "
            "re-enumerating the library from scratch.",
            drive_id,
        )
        delta_links.pop(drive_id, None)
        items, delta_link = _collect_delta(session, initial_url)

    yielded = 0
    tombstoned = 0
    for item in items.values():
        if "deleted" in item:
            tombstoned += 1
            yield {"id": _row_id(drive_id, item["id"]), "_deleted": True}
            continue

        row = _item_to_row(session, config, drive_id, item)
        if row is not None:
            yielded += 1
            yield row

    # Only advance once the whole feed has been emitted, so a failure mid-sync
    # replays from the same link instead of skipping changes.
    if delta_link:
        delta_links[drive_id] = delta_link
    logger.info(
        "SharePoint: drive %s synced %d file(s), %d deletion(s).", drive_id, yielded, tombstoned
    )


def _iter_rows(session: Any, config: _SharePointConfig, state: dict):
    """Yield one row per in-scope file across the configured libraries."""
    delta_links = state.setdefault("delta_links", {})
    for drive_id in _resolve_drive_ids(session, config):
        yield from _sync_drive(session, config, drive_id, delta_links)


# ---------------------------------------------------------------------------
# Public factory
# ---------------------------------------------------------------------------
def sharepoint_source(
    site_id: str | None = None,
    *,
    drive_ids: list[str] | None = None,
    tenant_id: str | None = None,
    client_id: str | None = None,
    client_secret: str | None = None,
    max_file_size_mb: int | None = None,
    session: Any = None,
):
    """Return a ``dlt`` resource yielding one row per in-scope SharePoint file.

    Any argument left as ``None`` falls back to the matching environment
    variable.  Hand the result to ``cognee.remember(...)`` with
    ``write_disposition="merge"`` and ``primary_key="id"``.

    Args:
        site_id: Graph site id whose document libraries are synced. Falls back
            to ``SHAREPOINT_SITE_ID``.
        drive_ids: Sync only these drive (document library) ids instead of every
            library on the site. Falls back to the comma-separated
            ``SHAREPOINT_DRIVE_IDS``.
        tenant_id: Entra ID tenant. Falls back to ``MICROSOFT_TENANT_ID``.
        client_id: Entra ID app registration id. Falls back to ``MICROSOFT_CLIENT_ID``.
        client_secret: Entra ID client secret. Falls back to ``MICROSOFT_CLIENT_SECRET``.
        max_file_size_mb: Skip files larger than this (default 25).
        session: Pre-built ``requests`` session. Mainly an injection point for
            tests; when omitted one is built from the credentials above.
    """
    try:
        import dlt
    except ImportError as exc:
        raise ImportError(
            'The SharePoint connector requires the dlt extra: pip install "cognee[dlt]".'
        ) from exc

    resolved_site_id = site_id or os.getenv("SHAREPOINT_SITE_ID")
    resolved_drive_ids = drive_ids or _split_ids(os.getenv("SHAREPOINT_DRIVE_IDS"))
    if not resolved_site_id and not resolved_drive_ids:
        raise ValueError(
            "sharepoint_source requires site_id or drive_ids (pass one explicitly "
            "or set SHAREPOINT_SITE_ID / SHAREPOINT_DRIVE_IDS)."
        )

    config = _SharePointConfig(
        site_id=resolved_site_id,
        drive_ids=tuple(resolved_drive_ids),
        max_file_size_mb=(
            max_file_size_mb
            if max_file_size_mb is not None
            else int(os.getenv("SHAREPOINT_MAX_FILE_SIZE_MB", "25"))
        ),
    )

    @dlt.resource(
        name="sharepoint_files",
        primary_key="id",
        write_disposition="merge",
        # _deleted is a boolean hard-delete marker: rows where it is True are
        # removed from the dlt destination on merge, which propagates the
        # deletion through cognee's orphan_cleanup.
        columns={"_deleted": {"data_type": "bool", "hard_delete": True}},
    )
    def sharepoint_files():
        client = session or build_graph_session(
            tenant_id=tenant_id, client_id=client_id, client_secret=client_secret
        )
        yield from _iter_rows(client, config, dlt.current.resource_state())

    resource = sharepoint_files()
    # Opt into the document ingestion path: each file row (id/title/content/url)
    # becomes a text document that flows through normal cognify.
    setattr(resource, DOCUMENT_SOURCE_ATTR, "sharepoint")
    return resource


def _split_ids(value: str | None) -> list[str]:
    return [part.strip() for part in (value or "").split(",") if part.strip()]
